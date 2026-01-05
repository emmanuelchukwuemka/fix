from flask import Blueprint, request, jsonify
from backend.extensions import db
from backend.models.user import User, UserRole
from backend.models.reward_code import RewardCode
from backend.models.transaction import Transaction, TransactionType, TransactionStatus
from backend.models.support_message import SupportMessage, MessageStatus, MessageType
from backend.models.task import Task, UserTask
from backend.utils.helpers import generate_reward_code, generate_batch_id
from backend.utils.emailer import Emailer
from backend.utils.admin_auth import admin_required
from flask_jwt_extended import jwt_required, get_jwt_identity
import csv
import io

admin_bp = Blueprint('admin', __name__)

@admin_bp.route('/codes/generate', methods=['POST'])
@admin_required
def generate_codes():
    try:
        current_user_id = int(get_jwt_identity())
        user = User.query.get(current_user_id)
        
        data = request.get_json()
        count = data.get('count', 100)
        point_value = float(data.get('point_value', 1.0))  # Use float to allow decimal values
        
        if count <= 0 or count > 10000:
            return jsonify({'message': 'Count must be between 1 and 10,000'}), 400
        
        if point_value <= 0:
            return jsonify({'message': 'Point value must be greater than 0'}), 400
        
        # Generate batch ID
        batch_id = generate_batch_id()
        
        # Generate codes
        codes = []
        for _ in range(count):
            code = generate_reward_code()
            reward_code = RewardCode(
                code=code,
                point_value=point_value,  # This can now be a float value
                batch_id=batch_id
            )
            db.session.add(reward_code)
            codes.append(code)
        
        db.session.commit()
        
        return jsonify({
            'message': f'Successfully generated {count} codes',
            'batch_id': batch_id,
            'codes': codes[:10],  # Return first 10 codes as sample
            'total_generated': count
        }), 201
        
    except Exception as e:
        return jsonify({'message': 'Failed to generate codes', 'error': str(e)}), 500

@admin_bp.route('/codes/export/<batch_id>', methods=['GET'])
@admin_required
def export_codes(batch_id):
    try:
        current_user_id = int(get_jwt_identity())
        user = User.query.get(current_user_id)
        
        # Get format parameter (csv, pdf, word)
        format_type = request.args.get('format', 'csv').lower()
        
        # Get search and status filters
        search = request.args.get('search', '')
        status = request.args.get('status', '')
        
        # Build query based on batch_id
        query = RewardCode.query
        
        if batch_id == 'all':
            # If batch_id is 'all', don't filter by batch_id
            pass
        else:
            # Filter by specific batch_id
            query = query.filter_by(batch_id=batch_id)
        
        # Apply search filter if provided
        if search:
            query = query.filter(RewardCode.code.ilike(f'%{search}%'))
        
        # Apply status filter if provided
        if status == 'available':
            query = query.filter_by(is_used=False)
        elif status == 'used':
            query = query.filter_by(is_used=True)
        
        # Get the codes
        codes = query.all()
        
        if not codes:
            if batch_id == 'all':
                return jsonify({'message': 'No codes found matching the criteria'}), 404
            else:
                return jsonify({'message': 'No codes found for this batch'}), 404
        
        if format_type == 'csv':
            # Create CSV in memory
            output = io.StringIO()
            writer = csv.writer(output)
            writer.writerow(['Code', 'Point Value', 'USD Value', 'Batch ID', 'Created At'])
            
            for code in codes:
                usd_value = code.point_value * 0.30  # Calculate USD value based on 1 point = $0.30
                writer.writerow([
                    code.code,
                    code.point_value,
                    f"{usd_value:.2f}",  # Format to 2 decimal places
                    code.batch_id,
                    code.created_at.isoformat()
                ])
            
            csv_data = output.getvalue()
            output.close()
            
            return jsonify({
                'batch_id': batch_id,
                'code_count': len(codes),
                'csv_data': csv_data,
                'format': 'csv'
            }), 200
        
        elif format_type == 'pdf':
            # Create PDF using reportlab
            from reportlab.lib.pagesizes import letter
            from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.lib import colors
            import base64
            
            # Create in-memory buffer
            buffer = io.BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=letter)
            elements = []
            
            # Add title
            styles = getSampleStyleSheet()
            title = Paragraph(f"MyFigPoint - Reward Codes Report (Batch: {batch_id})", styles['Title'])
            elements.append(title)
            elements.append(Spacer(1, 12))
            
            # Add subtitle with date
            from datetime import datetime
            date_str = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            subtitle = Paragraph(f"Generated on: {date_str}", styles['Normal'])
            elements.append(subtitle)
            elements.append(Spacer(1, 20))
            
            # Create table data
            data = [['Code', 'Point Value', 'USD Value', 'Batch ID', 'Created At']]
            for code in codes:
                usd_value = code.point_value * 0.30  # Calculate USD value based on 1 point = $0.30
                data.append([
                    code.code,
                    str(code.point_value),
                    f"{usd_value:.2f}",  # Format to 2 decimal places
                    code.batch_id or '',
                    code.created_at.isoformat()
                ])
            
            # Create table
            table = Table(data)
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 14),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 1), (-1, -1), 12),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            
            elements.append(table)
            
            # Build PDF
            doc.build(elements)
            
            # Get PDF data and encode it
            pdf_data = buffer.getvalue()
            buffer.close()
            encoded_pdf = base64.b64encode(pdf_data).decode('utf-8')
            
            return jsonify({
                'batch_id': batch_id,
                'code_count': len(codes),
                'pdf_data': encoded_pdf,
                'format': 'pdf'
            }), 200
        
        elif format_type == 'word':
            # Create Word document using python-docx
            from docx import Document
            from docx.shared import Inches
            import base64
            
            # Create in-memory buffer
            buffer = io.BytesIO()
            
            # Create Word document
            doc = Document()
            
            # Add title
            doc.add_heading(f'MyFigPoint - Reward Codes Report (Batch: {batch_id})', 0)
            
            # Add subtitle with date
            from datetime import datetime
            date_str = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            doc.add_paragraph(f'Generated on: {date_str}')
            
            # Add table
            table = doc.add_table(rows=1, cols=5)
            table.style = 'Table Grid'
            
            # Add header row
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Code'
            hdr_cells[1].text = 'Point Value'
            hdr_cells[2].text = 'USD Value'
            hdr_cells[3].text = 'Batch ID'
            hdr_cells[4].text = 'Created At'
            
            # Add data rows
            for code in codes:
                usd_value = code.point_value * 0.30  # Calculate USD value based on 1 point = $0.30
                row_cells = table.add_row().cells
                row_cells[0].text = code.code
                row_cells[1].text = str(code.point_value)
                row_cells[2].text = f"{usd_value:.2f}"  # Format to 2 decimal places
                row_cells[3].text = code.batch_id or ''
                row_cells[4].text = code.created_at.isoformat()
            
            # Save to buffer
            doc.save(buffer)
            
            # Get Word data and encode it
            word_data = buffer.getvalue()
            buffer.close()
            encoded_word = base64.b64encode(word_data).decode('utf-8')
            
            return jsonify({
                'batch_id': batch_id,
                'code_count': len(codes),
                'word_data': encoded_word,
                'format': 'word'
            }), 200
        
        else:
            return jsonify({'message': 'Invalid format. Use csv, pdf, or word'}), 400
        
    except Exception as e:
        return jsonify({'message': 'Failed to export codes', 'error': str(e)}), 500


@admin_bp.route('/codes/recent-batch', methods=['GET'])
@admin_required
def get_recent_batch():
    try:
        # Get the most recent batch ID based on the latest created code
        recent_code = RewardCode.query.order_by(RewardCode.created_at.desc()).first()
        
        if not recent_code or not recent_code.batch_id:
            return jsonify({'message': 'No batches found'}), 404
        
        return jsonify({
            'batch_id': recent_code.batch_id,
            'message': 'Successfully retrieved most recent batch'
        }), 200
        
    except Exception as e:
        return jsonify({'message': 'Failed to fetch recent batch', 'error': str(e)}), 500


@admin_bp.route('/codes', methods=['GET'])
@admin_required
def get_all_codes():
    try:
        current_user_id = int(get_jwt_identity())
        user = User.query.get(current_user_id)
        
        page = request.args.get('page', 1, type=int)
        per_page = request.args.get('per_page', 50, type=int)
        search = request.args.get('search', '')
        status = request.args.get('status', 'all') # all, available, used
        
        query = RewardCode.query
        
        if search:
            query = query.filter(RewardCode.code.ilike(f'%{search}%'))
            
        if status == 'available':
            query = query.filter_by(is_used=False)
        elif status == 'used':
            query = query.filter_by(is_used=True)
            
        codes = query.order_by(RewardCode.created_at.desc()).paginate(
            page=page, per_page=per_page, error_out=False
        )
        
        return jsonify({
            'codes': [code.to_dict() for code in codes.items],
            'total': codes.total,
            'pages': codes.pages,
            'current_page': page
        }), 200
        
    except Exception as e:
        return jsonify({'message': 'Failed to fetch codes', 'error': str(e)}), 500

@admin_bp.route('/codes/<int:code_id>/delete', methods=['DELETE'])
@admin_required
def delete_single_code(code_id):
    try:
        current_user_id = int(get_jwt_identity())
        user = User.query.get(current_user_id)
        
        code = RewardCode.query.get(code_id)
        if not code:
            return jsonify({'message': 'Code not found'}), 404
            
        db.session.delete(code)
        db.session.commit()
        
        return jsonify({'message': 'Code deleted successfully'}), 200
        
    except Exception as e:
        return jsonify({'message': 'Failed to delete code', 'error': str(e)}), 500


@admin_bp.route('/codes/<code>/details', methods=['GET'])
@admin_required
def get_code_details(code):
    try:
        current_user_id = int(get_jwt_identity())
        admin_user = User.query.get(current_user_id)
        
        # Find the reward code by its code value
        reward_code = RewardCode.query.filter_by(code=code.upper()).first()
        
        if not reward_code:
            return jsonify({'message': 'Code not found'}), 404
        
        # Prepare the code data
        code_data = reward_code.to_dict()
        
        # If the code was used, get user details
        if reward_code.used_by:
            user = User.query.get(reward_code.used_by)
            if user:
                code_data['used_by_user'] = {
                    'id': user.id,
                    'full_name': user.full_name,
                    'email': user.email
                }
        
        return jsonify({'code': code_data}), 200
        
    except Exception as e:
        return jsonify({'message': 'Failed to fetch code details', 'error': str(e)}), 500

@admin_bp.route('/support-messages', methods=['GET'])
@admin_required
def get_all_support_messages():
    try:
        current_user_id = int(get_jwt_identity())
        admin_user = User.query.get(current_user_id)
        
        page = request.args.get('page', 1, type=int)
        per_page = request.args.get('per_page', 20, type=int)
        status = request.args.get('status', 'all')
        
        query = SupportMessage.query
        
        if status != 'all':
            query = query.filter_by(status=MessageStatus(status))
            
        messages = query.order_by(SupportMessage.created_at.desc()).paginate(
            page=page, per_page=per_page, error_out=False
        )
        
        # Build user mapping to avoid N+1 queries
        user_ids = [msg.user_id for msg in messages.items]
        users = User.query.filter(User.id.in_(user_ids)).all()
        user_map = {user.id: user.to_dict() for user in users}
        
        result = []
        for msg in messages.items:
            msg_data = msg.to_dict()
            msg_data['user'] = user_map.get(msg.user_id)
            result.append(msg_data)
        
        return jsonify({
            'messages': result,
            'total': messages.total,
            'pages': messages.pages,
            'current_page': page
        }), 200
        
    except Exception as e:
        return jsonify({'message': 'Failed to fetch support messages', 'error': str(e)}), 500

@admin_bp.route('/support-messages/<int:message_id>/respond', methods=['POST'])
@admin_required
def respond_to_support_message(message_id):
    try:
        current_user_id = int(get_jwt_identity())
        admin_user = User.query.get(current_user_id)
        
        data = request.get_json()
        response_text = data.get('response')
        
        if not response_text:
            return jsonify({'message': 'Response text is required'}), 400
            
        message = SupportMessage.query.get(message_id)
        if not message:
            return jsonify({'message': 'Support message not found'}), 404
            
        message.response = response_text
        message.status = MessageStatus.REPLIED
        message.updated_at = db.func.current_timestamp()
        
        db.session.commit()
        
        return jsonify({
            'message': 'Response sent successfully',
            'support_message': message.to_dict()
        }), 200
        
    except Exception as e:
        return jsonify({'message': 'Failed to respond to support message', 'error': str(e)}), 500

@admin_bp.route('/activities/recent', methods=['GET'])
@admin_required
def get_recent_activities():
    try:
        current_user_id = int(get_jwt_identity())
        admin_user = User.query.get(current_user_id)
            
        limit = request.args.get('limit', 20, type=int)
        
        # Combine different types of activities
        activities = []
        
        # 1. New user registrations
        new_users = User.query.order_by(User.created_at.desc()).limit(limit).all()
        for user in new_users:
            activities.append({
                'type': 'user_registration',
                'user_name': user.full_name,
                'user_id': user.id,
                'timestamp': user.created_at.isoformat(),
                'referral_code': user.referral_code,
                'points': 0,
                'cash_value': 0
            })
            
        # 2. Recent transactions (Task completions, withdrawals, redemptions)
        transactions = Transaction.query.order_by(Transaction.created_at.desc()).limit(limit).all()
        for tx in transactions:
            user = tx.user
            act_type = 'unknown'
            if tx.type == TransactionType.EARNING:
                act_type = 'task_completion'
            elif tx.type == TransactionType.CODE_REDEMPTION:
                act_type = 'code_redemption'
            elif tx.type == TransactionType.POINT_WITHDRAWAL:
                act_type = 'withdrawal_request'
            elif tx.type == TransactionType.REFERRAL_BONUS:
                act_type = 'referral_bonus'
                
            activities.append({
                'type': act_type,
                'user_name': user.full_name if user else 'Unknown',
                'user_id': tx.user_id,
                'timestamp': tx.created_at.isoformat(),
                'points': abs(tx.points_amount or 0),
                'cash_value': abs(tx.amount or 0),
                'task_title': tx.description,
                'code': tx.reference_id if tx.type == TransactionType.CODE_REDEMPTION else None
            })
            
        # Sort combined activities by timestamp
        activities.sort(key=lambda x: x['timestamp'], reverse=True)
        
        return jsonify({
            'activities': activities[:limit]
        }), 200
        
    except Exception as e:
        return jsonify({'message': 'Failed to fetch recent activities', 'error': str(e)}), 500

@admin_bp.route('/users/points/update', methods=['POST'])
@admin_required
def update_user_points():
    try:
        current_user_id = int(get_jwt_identity())
        admin_user = User.query.get(current_user_id)
        
        data = request.get_json()
        user_id = data.get('user_id')
        points = int(data.get('points', 0))  # Ensure points is an integer
        operation = data.get('operation', 'add')  # add or subtract
        
        if not user_id:
            return jsonify({'message': 'User ID is required'}), 400
        
        user = User.query.get(user_id)
        if not user:
            return jsonify({'message': 'User not found'}), 404
        
        # Update points based on operation
        if operation == 'add':
            raw_points = int(points)  # Ensure points is an integer
            points_to_add = max(1, raw_points)  # Ensure at least 1 point if any points are added
            user.points_balance += points_to_add
            user.total_points_earned += points_to_add
        elif operation == 'subtract':
            points_to_subtract = max(1, int(points))  # Ensure at least 1 point if any points are subtracted
            if user.points_balance < points_to_subtract:
                return jsonify({'message': 'Insufficient points balance'}), 400
            user.points_balance -= points_to_subtract
            user.total_points_withdrawn += points_to_subtract
        else:
            return jsonify({'message': 'Invalid operation. Use "add" or "subtract"'}), 400
        
        db.session.commit()
        
        return jsonify({
            'message': f'Successfully {"added" if operation == "add" else "subtracted"} {points} points',
            'user_id': user_id,
            'new_balance': user.points_balance
        }), 200
        
    except Exception as e:
        return jsonify({'message': 'Failed to update user points', 'error': str(e)}), 500

@admin_bp.route('/users', methods=['GET'])
@admin_required
def get_all_users():
    try:
        current_user_id = int(get_jwt_identity())
        admin_user = User.query.get(current_user_id)
        
        page = request.args.get('page', 1, type=int)
        per_page = request.args.get('per_page', 20, type=int)
        search = request.args.get('search', '')
        role_filter = request.args.get('role', '')
        status_filter = request.args.get('status', '')
        
        query = User.query
        
        if search:
            query = query.filter(
                db.or_(
                    User.full_name.ilike(f'%{search}%'),
                    User.email.ilike(f'%{search}%')
                )
            )
        
        # Filter by role if provided
        if role_filter:
            query = query.filter(User.role == UserRole(role_filter))
        
        # Filter by status if provided
        if status_filter:
            if status_filter == 'active':
                query = query.filter(User.is_suspended == False)
            elif status_filter == 'suspended':
                query = query.filter(User.is_suspended == True)
        
        users = query.order_by(User.created_at.desc()).paginate(
            page=page, per_page=per_page, error_out=False
        )
        
        return jsonify({
            'users': [user.to_dict() for user in users.items],
            'total': users.total,
            'pages': users.pages,
            'current_page': page
        }), 200
        
    except Exception as e:
        return jsonify({'message': 'Failed to fetch users', 'error': str(e)}), 500

@admin_bp.route('/users/<int:user_id>/details', methods=['GET'])
@admin_required
def get_user_details(user_id):
    try:
        current_user_id = int(get_jwt_identity())
        admin_user = User.query.get(current_user_id)
        
        user = User.query.get(user_id)
        if not user:
            return jsonify({'message': 'User not found'}), 404
            
        return jsonify({'user': user.to_dict()}), 200
        
    except Exception as e:
        return jsonify({'message': 'Failed to fetch user details', 'error': str(e)}), 500

@admin_bp.route('/users/<int:user_id>/suspend', methods=['POST'])
@admin_required
def suspend_user(user_id):
    try:
        current_user_id = int(get_jwt_identity())
        admin_user = User.query.get(current_user_id)
            
        user = User.query.get(user_id)
        if not user:
            return jsonify({'message': 'User not found'}), 404
            
        user.is_suspended = True
        db.session.commit()
        
        return jsonify({'message': 'User suspended successfully'}), 200
    except Exception as e:
        return jsonify({'message': 'Failed to suspend user', 'error': str(e)}), 500

@admin_bp.route('/users/<int:user_id>/unsuspend', methods=['POST'])
@admin_required
def unsuspend_user(user_id):
    try:
        current_user_id = int(get_jwt_identity())
        admin_user = User.query.get(current_user_id)
            
        user = User.query.get(user_id)
        if not user:
            return jsonify({'message': 'User not found'}), 404
            
        user.is_suspended = False
        db.session.commit()
        
        return jsonify({'message': 'User unsuspended successfully'}), 200
    except Exception as e:
        return jsonify({'message': 'Failed to unsuspend user', 'error': str(e)}), 500

@admin_bp.route('/users/<int:user_id>/verify', methods=['POST'])
@admin_required
def verify_user(user_id):
    try:
        current_user_id = int(get_jwt_identity())
        admin_user = User.query.get(current_user_id)
            
        user = User.query.get(user_id)
        if not user:
            return jsonify({'message': 'User not found'}), 404
            
        data = request.get_json()
        is_verified = data.get('verified', True)
        
        user.is_verified = is_verified
        user.verification_pending = False
        db.session.commit()
        
        return jsonify({'message': f'User verification status updated to {is_verified}'}), 200
    except Exception as e:
        return jsonify({'message': 'Failed to verify user', 'error': str(e)}), 500


@admin_bp.route('/users/<int:user_id>/update', methods=['POST'])
@admin_required
def update_user(user_id):
    try:
        current_user_id = int(get_jwt_identity())
        admin_user = User.query.get(current_user_id)
        
        user = User.query.get(user_id)
        if not user:
            return jsonify({'message': 'User not found'}), 404
        
        data = request.get_json()
        
        # Update user fields
        user.full_name = data.get('full_name', user.full_name)
        user.email = data.get('email', user.email)
        user.phone = data.get('phone', user.phone)
        user.role = UserRole(data.get('role', user.role.value if hasattr(user.role, 'value') else user.role))
        user.bank_name = data.get('bank_name', user.bank_name)
        user.account_name = data.get('account_name', user.account_name)
        user.account_number = data.get('account_number', user.account_number)
        user.routing_number = data.get('routing_number', user.routing_number)
        user.points_balance = float(data.get('points_balance', user.points_balance))
        user.total_earnings = float(data.get('total_earnings', user.total_earnings))
        user.is_verified = data.get('is_verified', user.is_verified)
        user.is_suspended = data.get('is_suspended', user.is_suspended)
        user.is_approved = data.get('is_approved', user.is_approved)
        
        db.session.commit()
        
        return jsonify({
            'message': 'User updated successfully',
            'user': user.to_dict()
        }), 200
    
    except Exception as e:
        return jsonify({'message': 'Failed to update user', 'error': str(e)}), 500


@admin_bp.route('/users/<int:user_id>/verify-documents', methods=['POST'])
@admin_required
def verify_user_documents(user_id):
    try:
        current_user_id = int(get_jwt_identity())
        admin_user = User.query.get(current_user_id)
        
        user = User.query.get(user_id)
        if not user:
            return jsonify({'message': 'User not found'}), 404
        
        # Mark user as verified
        user.is_verified = True
        user.verification_pending = False
        
        # Create notification for the user
        from backend.models.notification import Notification, NotificationType
        notification = Notification(
            user_id=user.id,
            title="Documents Verified",
            message="Your submitted documents have been verified successfully by an administrator.",
            type=NotificationType.SUCCESS
        )
        
        db.session.add(notification)
        db.session.commit()
        
        # Optionally send email notification
        try:
            from backend.utils.emailer import send_email
            user_email = user.email
            email_subject = "Document Verification Approved - MyFigPoint"
            email_body = f"Hello {user.full_name},\n\nYour submitted documents have been verified successfully by an administrator. Your account is now verified.\n\nThank you for using MyFigPoint!"
            send_email(user_email, email_subject, email_body)
        except Exception as email_error:
            # If email fails, log the error but don't fail the entire operation
            print(f"Failed to send email notification: {str(email_error)}")
        
        return jsonify({
            'message': 'Documents verified successfully',
            'user': user.to_dict()
        }), 200
        
    except Exception as e:
        return jsonify({'message': 'Failed to verify documents', 'error': str(e)}), 500


@admin_bp.route('/users/<int:user_id>/message', methods=['POST'])
@admin_required
def send_user_message(user_id):
    try:
        current_user_id = int(get_jwt_identity())
        admin_user = User.query.get(current_user_id)
        
        user = User.query.get(user_id)
        if not user:
            return jsonify({'message': 'User not found'}), 404
        
        data = request.get_json()
        message_text = data.get('message', '')
        
        if not message_text:
            return jsonify({'message': 'Message content is required'}), 400
        
        # Create notification for the user
        from backend.models.notification import Notification, NotificationType
        notification = Notification(
            user_id=user.id,
            title="Message from Admin",
            message=message_text,
            type=NotificationType.INFO
        )
        
        db.session.add(notification)
        db.session.commit()
        
        # Optionally send email notification
        try:
            from backend.utils.emailer import send_email
            user_email = user.email
            email_subject = "Message from Admin - MyFigPoint"
            email_body = f"Hello {user.full_name},\n\nYou have received a message from an administrator:\n\n{message_text}\n\nThank you for using MyFigPoint!"
            send_email(user_email, email_subject, email_body)
        except Exception as email_error:
            # If email fails, log the error but don't fail the entire operation
            print(f"Failed to send email notification: {str(email_error)}")
        
        return jsonify({
            'message': 'Message sent successfully',
            'user': user.to_dict()
        }), 200
        
    except Exception as e:
        return jsonify({'message': 'Failed to send message', 'error': str(e)}), 500

@admin_bp.route('/withdrawals', methods=['GET'])
@admin_required
def get_pending_withdrawals():
    try:
        current_user_id = int(get_jwt_identity())
        admin_user = User.query.get(current_user_id)
        
        page = request.args.get('page', 1, type=int)
        per_page = request.args.get('per_page', 20, type=int)
        status = request.args.get('status', 'pending')  # Default to pending withdrawals
        
        # Query for withdrawal transactions
        query = Transaction.query.filter_by(type=TransactionType.POINT_WITHDRAWAL)
        
        if status:
            query = query.filter_by(status=TransactionStatus(status))
        
        withdrawals = query.order_by(Transaction.created_at.desc()).paginate(
            page=page, per_page=per_page, error_out=False
        )
        
        # Get user details for each withdrawal
        withdrawal_data = []
        for withdrawal in withdrawals.items:
            user = User.query.get(withdrawal.user_id)
            withdrawal_data.append({
                'transaction': withdrawal.to_dict(),
                'user': user.to_dict() if user else None
            })
        
        # Calculate stats for the overview
        from datetime import datetime, timedelta
        today = datetime.utcnow().replace(hour=0, minute=0, second=0, microsecond=0)
        week_ago = today - timedelta(days=7)
        
        pending_count = Transaction.query.filter_by(
            type=TransactionType.POINT_WITHDRAWAL,
            status=TransactionStatus.PENDING
        ).count()
        
        today_approved = Transaction.query.filter(
            Transaction.type == TransactionType.POINT_WITHDRAWAL,
            Transaction.status == TransactionStatus.COMPLETED,
            Transaction.updated_at >= today
        ).count()
        
        week_approved_sum = db.session.query(db.func.coalesce(db.func.sum(Transaction.amount), 0.0)).filter(
            Transaction.type == TransactionType.POINT_WITHDRAWAL,
            Transaction.status == TransactionStatus.COMPLETED,
            Transaction.updated_at >= week_ago
        ).scalar() or 0.0
        
        rejected_count = Transaction.query.filter(
            Transaction.type == TransactionType.POINT_WITHDRAWAL,
            Transaction.status == TransactionStatus.FAILED
        ).count()
        
        return jsonify({
            'withdrawals': withdrawal_data,
            'total': withdrawals.total,
            'pages': withdrawals.pages,
            'current_page': page,
            'stats': {
                'pending': pending_count,
                'today_approved': today_approved,
                'week_total_amount': abs(week_approved_sum),
                'rejected': rejected_count
            }
        }), 200
        
    except Exception as e:
        return jsonify({'message': 'Failed to fetch withdrawals', 'error': str(e)}), 500

@admin_bp.route('/withdrawals/<int:transaction_id>/approve', methods=['POST'])
@admin_required
def approve_withdrawal(transaction_id):
    try:
        current_user_id = int(get_jwt_identity())
        admin_user = User.query.get(current_user_id)
        
        # Get the transaction
        transaction = Transaction.query.get(transaction_id)
        if not transaction:
            return jsonify({'message': 'Transaction not found'}), 404
        
        if transaction.type != TransactionType.POINT_WITHDRAWAL:
            return jsonify({'message': 'Transaction is not a withdrawal request'}), 400
        
        if transaction.status != TransactionStatus.PENDING:
            return jsonify({'message': 'Transaction is not pending approval'}), 400
        
        # Get the user
        user = User.query.get(transaction.user_id)
        if not user:
            return jsonify({'message': 'User not found'}), 404
        
        # Verify that points were already deducted when the withdrawal was requested
        # The points should have been deducted when the request was submitted
        points_to_deduct = abs(transaction.points_amount)
        
        # Update transaction status to completed
        transaction.status = TransactionStatus.COMPLETED
        transaction.updated_at = db.func.current_timestamp()
        
        db.session.commit()
        
        # Send email notification
        emailer = Emailer()
        emailer.send_withdrawal_approved_notification(
            user_email=user.email,
            user_name=user.full_name,
            points=abs(transaction.points_amount),
            amount=abs(transaction.amount),
            method=get_method_from_description(transaction.description)
        )
        
        return jsonify({
            'message': 'Withdrawal approved successfully',
            'transaction': transaction.to_dict(),
            'new_balance': user.points_balance,
            'new_usd_balance': user.total_earnings  # This reflects the USD equivalent
        }), 200
        
    except Exception as e:
        return jsonify({'message': 'Failed to approve withdrawal', 'error': str(e)}), 500

@admin_bp.route('/withdrawals/<int:transaction_id>/reject', methods=['POST'])
@admin_required
def reject_withdrawal(transaction_id):
    try:
        current_user_id = int(get_jwt_identity())
        admin_user = User.query.get(current_user_id)
        
        # Get the transaction
        transaction = Transaction.query.get(transaction_id)
        if not transaction:
            return jsonify({'message': 'Transaction not found'}), 404
        
        if transaction.type != TransactionType.POINT_WITHDRAWAL:
            return jsonify({'message': 'Transaction is not a withdrawal request'}), 400
        
        if transaction.status != TransactionStatus.PENDING:
            return jsonify({'message': 'Transaction is not pending approval'}), 400
        
        # Get the user
        user = User.query.get(transaction.user_id)
        if not user:
            return jsonify({'message': 'User not found'}), 404
        
        # Refund points to user
        # points_amount is negative for withdrawals, so subtracting a negative value adds the points back
        user.points_balance -= transaction.points_amount
        user.total_points_withdrawn += transaction.points_amount  # This will reduce the total withdrawn since points_amount is negative
        user.total_withdrawn += transaction.amount  # This will reduce the total withdrawn since amount is negative
        
        # Update transaction status to failed
        transaction.status = TransactionStatus.FAILED
        transaction.updated_at = db.func.current_timestamp()
        
        db.session.commit()
        
        # Send email notification
        emailer = Emailer()
        emailer.send_withdrawal_rejected_notification(
            user_email=user.email,
            user_name=user.full_name,
            points=abs(transaction.points_amount),
            amount=abs(transaction.amount),
            method=get_method_from_description(transaction.description),
            reason="Administrative review"
        )
        
        return jsonify({
            'message': 'Withdrawal rejected and points refunded',
            'transaction': transaction.to_dict(),
            'new_balance': user.points_balance,
            'new_usd_balance': user.total_earnings  # This reflects the USD equivalent
        }), 200
        
    except Exception as e:
        return jsonify({'message': 'Failed to reject withdrawal', 'error': str(e)}), 500



def get_method_from_description(description):
    """Extract payment method from transaction description"""
    import re
    match = re.search(r'via (\w+)', description)
    return match.group(1) if match else 'unknown'

@admin_bp.route('/referrals/award-bonus', methods=['POST'])
@admin_required
def award_referral_bonus():
    try:
        current_user_id = int(get_jwt_identity())
        admin_user = User.query.get(current_user_id)
        
        data = request.get_json()
        user_id = data.get('user_id')
        points = int(data.get('points', 0))  # Ensure points is an integer
        amount = float(data.get('amount', 0.0))  # Amount can be decimal
        reason = data.get('reason', 'Admin awarded referral bonus')
        
        if not user_id:
            return jsonify({'message': 'User ID is required'}), 400
        
        user = User.query.get(user_id)
        if not user:
            return jsonify({'message': 'User not found'}), 404
        
        # Award points/amount to user
        if points > 0:
            raw_points = int(points)  # Ensure points is an integer
            points_to_add = max(1, raw_points)  # Ensure at least 1 point if any points are awarded
            user.points_balance += points_to_add
            user.total_points_earned += points_to_add
        else:
            points_to_add = 0  # No points to add
            
        if amount > 0:
            user.total_earnings += amount
        
        # Create referral bonus transaction
        transaction = Transaction(
            user_id=user_id,
            type=TransactionType.REFERRAL_BONUS,
            status=TransactionStatus.COMPLETED,
            description=reason,
            amount=amount,
            points_amount=points_to_add if points > 0 else 0
        )
        
        db.session.add(transaction)
        db.session.commit()
        
        return jsonify({
            'message': f'Successfully awarded referral bonus to {user.full_name}',
            'user_id': user_id,
            'new_balance': user.points_balance
        }), 200
        
    except Exception as e:
        return jsonify({'message': 'Failed to award referral bonus', 'error': str(e)}), 500

@admin_bp.route('/dashboard/stats', methods=['GET'])
@admin_required
def get_dashboard_stats():
    try:
        current_user_id = int(get_jwt_identity())
        admin_user = User.query.get(current_user_id)
        
        # Get total users
        total_users = User.query.count()
        
        # Get pending withdrawals
        pending_withdrawals = Transaction.query.filter_by(
            type=TransactionType.POINT_WITHDRAWAL,
            status=TransactionStatus.PENDING
        ).count()
        
        # Get active tasks
        from backend.models.task import Task
        active_tasks = Task.query.count()
        
        # Get pending support messages
        pending_support = SupportMessage.query.filter_by(
            status=MessageStatus.SENT
        ).count()
        
        # Get platform earnings (sum of completed earning transactions)
        platform_earnings = db.session.query(db.func.coalesce(db.func.sum(Transaction.amount), 0.0)).filter(
            Transaction.status == TransactionStatus.COMPLETED,
            Transaction.type != TransactionType.POINT_WITHDRAWAL  # Exclude withdrawal transactions from platform earnings
        ).scalar() or 0.0
        
        # Get daily active users (placeholder)
        from datetime import datetime, timedelta
        yesterday = datetime.utcnow() - timedelta(days=1)
        daily_active_users = User.query.filter(User.updated_at >= yesterday).count()
        
        # Get tasks completed today
        from backend.models.task import UserTask
        from datetime import datetime, date
        today_start = datetime.combine(date.today(), datetime.min.time())
        tasks_completed_today = UserTask.query.filter(
            UserTask.status == 'completed',
            UserTask.updated_at >= today_start
        ).count()
        
        # Get total referrals
        total_referrals = User.query.filter(User.referred_by.isnot(None)).count()
        
        # Get referral earnings
        referral_earnings = db.session.query(db.func.coalesce(db.func.sum(Transaction.amount), 0.0)).filter(
            Transaction.type == TransactionType.REFERRAL_BONUS,
            Transaction.status == TransactionStatus.COMPLETED
        ).scalar() or 0.0
        
        # Get partners count
        total_partners = User.query.filter_by(role=UserRole.PARTNER).count()
        pending_approvals = User.query.filter_by(role=UserRole.PARTNER, is_approved=False).count()
        approved_partners = User.query.filter_by(role=UserRole.PARTNER, is_approved=True).count()
        
        return jsonify({
            'stats': {
                'total_users': total_users,
                'pending_withdrawals': pending_withdrawals,
                'active_tasks': active_tasks,
                'pending_support': pending_support,
                'platform_earnings': abs(platform_earnings),
                'daily_active_users': daily_active_users,
                'tasks_completed_today': tasks_completed_today,
                'total_referrals': total_referrals,
                'referral_earnings': abs(referral_earnings),
                'total_partners': total_partners,
                'pending_approvals': pending_approvals,
                'approved_partners': approved_partners
            }
        }), 200
        
    except Exception as e:
        return jsonify({'message': 'Failed to fetch dashboard stats', 'error': str(e)}), 500

@admin_bp.route('/tasks', methods=['GET'])
@admin_required
def get_all_tasks():
    try:
        current_user_id = int(get_jwt_identity())
        admin_user = User.query.get(current_user_id)
        
        page = request.args.get('page', 1, type=int)
        per_page = request.args.get('per_page', 20, type=int)
        search = request.args.get('search', '')
        status = request.args.get('status', 'all')  # all, active, inactive
        
        query = Task.query
        
        if search:
            query = query.filter(
                db.or_(
                    Task.title.ilike(f'%{search}%'),
                    Task.description.ilike(f'%{search}%')
                )
            )
        
        if status == 'active':
            query = query.filter_by(is_active=True)
        elif status == 'inactive':
            query = query.filter_by(is_active=False)
        
        tasks = query.order_by(Task.created_at.desc()).paginate(
            page=page, per_page=per_page, error_out=False
        )
        
        return jsonify({
            'tasks': [task.to_dict() for task in tasks.items],
            'total': tasks.total,
            'pages': tasks.pages,
            'current_page': page
        }), 200
        
    except Exception as e:
        return jsonify({'message': 'Failed to fetch tasks', 'error': str(e)}), 500

@admin_bp.route('/tasks/<int:task_id>/details', methods=['GET'])
@admin_required
def get_task_details(task_id):
    try:
        current_user_id = int(get_jwt_identity())
        admin_user = User.query.get(current_user_id)
        
        task = Task.query.get(task_id)
        if not task:
            return jsonify({'message': 'Task not found'}), 404
            
        return jsonify({'task': task.to_dict()}), 200
        
    except Exception as e:
        return jsonify({'message': 'Failed to fetch task details', 'error': str(e)}), 500

@admin_bp.route('/tasks', methods=['POST'])
@admin_required
def create_task():
    try:
        current_user_id = int(get_jwt_identity())
        admin_user = User.query.get(current_user_id)
        
        data = request.get_json()
        
        # Ensure is_active is set to True by default if not provided
        is_active = data.get('is_active', True)
        
        task = Task(
            title=data.get('title'),
            description=data.get('description', ''),
            reward_amount=data.get('reward_amount', 0.0),
            points_reward=int(data.get('points_reward', 0)),  # Ensure points is an integer
            category=data.get('category', 'General'),
            time_required=data.get('time_required', 0),
            is_active=is_active,
            requires_admin_verification=data.get('requires_admin_verification', False)  # New field
        )
        
        db.session.add(task)
        db.session.commit()
        
        # Refresh the task object to ensure all data is current
        db.session.refresh(task)
        
        print(f"Task created successfully: {task.id} - {task.title}, is_active: {task.is_active}")  # Debug print
        
        return jsonify({
            'message': 'Task created successfully',
            'task': task.to_dict()
        }), 201
        
    except Exception as e:
        db.session.rollback()  # Rollback in case of error
        print(f"Error creating task: {str(e)}")  # Debug print
        return jsonify({'message': 'Failed to create task', 'error': str(e)}), 500

@admin_bp.route('/tasks/<int:task_id>/update', methods=['POST'])
@admin_required
def update_task(task_id):
    try:
        current_user_id = int(get_jwt_identity())
        admin_user = User.query.get(current_user_id)
        
        data = request.get_json()
        task = Task.query.get(task_id)
        if not task:
            return jsonify({'message': 'Task not found'}), 404
        
        task.title = data.get('title', task.title)
        task.description = data.get('description', task.description)
        task.reward_amount = data.get('reward_amount', task.reward_amount)
        task.points_reward = int(data.get('points_reward', task.points_reward))  # Ensure points is an integer
        task.category = data.get('category', task.category)
        task.time_required = data.get('time_required', task.time_required)
        task.is_active = data.get('is_active', task.is_active)
        task.requires_admin_verification = data.get('requires_admin_verification', task.requires_admin_verification)  # New field
        
        db.session.commit()
        
        # Refresh the task object to ensure all data is current
        db.session.refresh(task)
        
        print(f"Task updated successfully: {task.id} - {task.title}, is_active: {task.is_active}")  # Debug print
        
        return jsonify({
            'message': 'Task updated successfully',
            'task': task.to_dict()
        }), 200
        
    except Exception as e:
        db.session.rollback()  # Rollback in case of error
        print(f"Error updating task: {str(e)}")  # Debug print
        return jsonify({'message': 'Failed to update task', 'error': str(e)}), 500

@admin_bp.route('/tasks/<int:task_id>/delete', methods=['DELETE'])
@admin_required
def delete_task(task_id):
    try:
        current_user_id = int(get_jwt_identity())
        admin_user = User.query.get(current_user_id)
        
        task = Task.query.get(task_id)
        if not task:
            return jsonify({'message': 'Task not found'}), 404
        
        db.session.delete(task)
        db.session.commit()
        
        print(f"Task deleted successfully: {task_id}")  # Debug print
        
        return jsonify({'message': 'Task deleted successfully'}), 200
        
    except Exception as e:
        db.session.rollback()  # Rollback in case of error
        print(f"Error deleting task: {str(e)}")  # Debug print
        return jsonify({'message': 'Failed to delete task', 'error': str(e)}), 500
